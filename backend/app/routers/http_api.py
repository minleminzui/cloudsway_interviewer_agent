from __future__ import annotations

import io
from typing import List

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from ..database import get_session
from ..models import Note, Session, Turn
from ..schemas import ExportRequest, PlanResponse, SessionCreate, SessionCreateResponse, SessionSchema
from ..services.outline import outline_builder

router = APIRouter()


@router.post("/sessions", response_model=SessionCreateResponse)
async def create_session(payload: SessionCreate, db: AsyncSession = Depends(get_session)) -> SessionCreateResponse:
    session = Session(topic=payload.topic, interviewer=payload.interviewer, interviewee=payload.interviewee)
    db.add(session)
    await db.commit()
    await db.refresh(session)
    outline = await outline_builder.build(payload.topic)
    return SessionCreateResponse(session=session, outline=outline)


@router.get("/sessions", response_model=List[SessionSchema])
async def list_sessions(db: AsyncSession = Depends(get_session)) -> List[SessionSchema]:
    result = await db.execute(select(Session))
    return list(result.scalars())


@router.get("/sessions/{session_id}", response_model=SessionSchema)
async def get_session_detail(session_id: int, db: AsyncSession = Depends(get_session)) -> SessionSchema:
    session = await db.get(Session, session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    return session


@router.post("/plan", response_model=PlanResponse)
async def generate_plan(payload: SessionCreate) -> PlanResponse:
    return await outline_builder.build(payload.topic)


@router.post("/export")
async def export_summary(payload: ExportRequest, db: AsyncSession = Depends(get_session)) -> StreamingResponse:
    session = await db.get(Session, payload.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    turns = await db.execute(select(Turn).where(Turn.session_id == payload.session_id))
    notes = await db.execute(select(Note).where(Note.session_id == payload.session_id))
    turn_rows = list(turns.scalars())
    note_rows = list(notes.scalars())
    if payload.format == "docx":
        from docx import Document

        document = Document()
        document.add_heading(f"采访纪要 - {session.topic}", level=1)
        for turn in turn_rows:
            document.add_heading(f"{turn.speaker} ({turn.stage})", level=2)
            document.add_paragraph(turn.transcript)
        document.add_heading("要点", level=2)
        for note in note_rows:
            document.add_paragraph(f"[{note.category}] {note.content}")
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        filename = f"session-{session.id}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    if payload.format == "xlsx":
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Turns"
        ws.append(["Speaker", "Stage", "Transcript"])
        for turn in turn_rows:
            ws.append([turn.speaker, turn.stage, turn.transcript])
        ws_notes = wb.create_sheet("Notes")
        ws_notes.append(["Category", "Content", "Confidence", "Need Clarify"])
        for note in note_rows:
            ws_notes.append([note.category, note.content, note.confidence, note.requires_clarification])
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        filename = f"session-{session.id}.xlsx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    raise HTTPException(status_code=400, detail="Unsupported format")
